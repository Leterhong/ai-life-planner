"""File Parser Service - Parse PDF, DOCX, TXT, Markdown files."""
import os
from pathlib import Path
from typing import Optional

import fitz  # PyMuPDF
from docx import Document as DocxDocument
import markdown


class FileParserService:
    """Service for extracting text from various file formats."""

    @staticmethod
    def parse_pdf(file_path: str) -> str:
        """Extract text from PDF file using PyMuPDF."""
        text_parts = []
        try:
            doc = fitz.open(file_path)
            for page_num in range(len(doc)):
                page = doc[page_num]
                text_parts.append(page.get_text())
            doc.close()
            return "\n\n".join(text_parts)
        except Exception as e:
            raise Exception(f"PDF parsing failed: {str(e)}")

    @staticmethod
    def parse_docx(file_path: str) -> str:
        """Extract text from DOCX file using python-docx."""
        text_parts = []
        try:
            doc = DocxDocument(file_path)
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # Also extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))

            return "\n\n".join(text_parts)
        except Exception as e:
            raise Exception(f"DOCX parsing failed: {str(e)}")

    @staticmethod
    def parse_txt(file_path: str) -> str:
        """Extract text from TXT file."""
        try:
            encodings = ["utf-8", "gbk", "gb2312", "latin-1"]
            for encoding in encodings:
                try:
                    with open(file_path, "r", encoding=encoding) as f:
                        return f.read()
                except UnicodeDecodeError:
                    continue
            # If all encodings fail, read as binary and decode with errors ignored
            with open(file_path, "rb") as f:
                return f.read().decode("utf-8", errors="ignore")
        except Exception as e:
            raise Exception(f"TXT parsing failed: {str(e)}")

    @staticmethod
    def parse_markdown(file_path: str) -> str:
        """Extract text from Markdown file (keep original text, strip HTML)."""
        try:
            encodings = ["utf-8", "gbk", "gb2312"]
            content = None
            for encoding in encodings:
                try:
                    with open(file_path, "r", encoding=encoding) as f:
                        content = f.read()
                    break
                except UnicodeDecodeError:
                    continue

            if content is None:
                with open(file_path, "rb") as f:
                    content = f.read().decode("utf-8", errors="ignore")

            # Convert markdown to text (simplified - strip markdown syntax)
            html = markdown.markdown(content)
            # Simple HTML tag removal
            import re
            text = re.sub(r"<[^>]+>", "", html)
            return text
        except Exception as e:
            raise Exception(f"Markdown parsing failed: {str(e)}")

    def parse_file(self, file_path: str, file_type: Optional[str] = None) -> str:
        """
        Parse a file and extract text content.

        Args:
            file_path: Path to the file
            file_type: File type override (pdf, docx, txt, md)

        Returns:
            Extracted text content
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        if file_type is None:
            ext = Path(file_path).suffix.lower().lstrip(".")
            file_type = ext

        file_type = file_type.lower()

        parsers = {
            "pdf": self.parse_pdf,
            "docx": self.parse_docx,
            "doc": self.parse_docx,
            "txt": self.parse_txt,
            "md": self.parse_markdown,
            "markdown": self.parse_markdown,
        }

        parser = parsers.get(file_type)
        if not parser:
            raise ValueError(f"Unsupported file type: {file_type}")

        return parser(file_path)


# Singleton instance
file_parser = FileParserService()
